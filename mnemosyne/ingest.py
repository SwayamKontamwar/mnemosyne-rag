from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterable

from .models import Chunk, SourceDocument

SUPPORTED = {".md", ".markdown", ".txt", ".pdf", ".docx", ".pptx", ".csv", ".tsv", ".xlsx"}
LINK_PATTERN = re.compile(r"\[\[([^\]]+)\]\]")
TAG_PATTERN = re.compile(r"(?:^|\s)#([a-zA-Z0-9_\-/]+)")
FRONTMATTER_TAGS = re.compile(r"(?ms)^---\s*\n(.*?)\n---\s*\n?")


def discover(path: Path) -> Iterable[Path]:
    if path.is_file():
        if path.suffix.lower() in SUPPORTED:
            yield path
        return
    for candidate in sorted(path.rglob("*")):
        if candidate.is_file() and candidate.suffix.lower() in SUPPORTED:
            yield candidate


def file_digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def parse(path: Path) -> list[SourceDocument]:
    absolute = str(path.resolve())
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(path)
        documents: list[SourceDocument] = []
        for index, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if not text.strip():
                text = _ocr_pdf_page(path, index + 1)
            if text.strip():
                documents.append(
                    SourceDocument(
                        absolute,
                        path.stem,
                        text,
                        index + 1,
                        file_type="pdf",
                        folder=_folder_label(path),
                    )
                )
        return documents
    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(path)
            blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
            for table in document.tables:
                blocks.extend("\t".join(cell.text.strip() for cell in row.cells) for row in table.rows)
            text = "\n\n".join(blocks)
        except ImportError:
            from zipfile import ZipFile

            with ZipFile(path) as archive:
                root = ET.fromstring(archive.read("word/document.xml"))
            paragraphs = ["".join(node.text or "" for node in paragraph.iter() if node.tag.endswith("}t")) for paragraph in root.iter() if paragraph.tag.endswith("}p")]
            text = "\n\n".join(item for item in paragraphs if item.strip())
        return [_source_document(path, absolute, text, "docx")]
    if suffix == ".pptx":
        from zipfile import ZipFile

        documents: list[SourceDocument] = []
        with ZipFile(path) as archive:
            names = [name for name in archive.namelist() if name.startswith("ppt/slides/slide") and name.endswith(".xml")]
            names.sort(key=lambda name: int(re.search(r"slide(\d+)\.xml", name).group(1)))
            for index, name in enumerate(names, 1):
                if name.startswith("ppt/slides/slide") and name.endswith(".xml"):
                    root = ET.fromstring(archive.read(name))
                    text = "\n".join(node.text or "" for node in root.iter() if node.tag.endswith("}t"))
                    if text.strip():
                        base = _source_document(path, absolute, text, "pptx")
                        documents.append(SourceDocument(**(base.__dict__ | {"page": index})))
        return documents
    if suffix in {".csv", ".tsv"}:
        return [_source_document(path, absolute, path.read_text(encoding="utf-8", errors="ignore"), suffix.lstrip("."))]
    if suffix == ".xlsx":
        return [_source_document(path, absolute, _extract_xlsx_text(path), "xlsx")]
    return [_source_document(path, absolute, path.read_text(encoding="utf-8"), suffix.lstrip(".") or "txt")]


def chunk_document(doc: SourceDocument, size: int, overlap: int) -> list[Chunk]:
    lines = doc.text.splitlines()
    chunks: list[Chunk] = []
    cursor = 0
    ordinal = 0
    while cursor < len(lines):
        selected: list[str] = []
        chars = 0
        end = cursor
        while end < len(lines) and (chars < size or not selected):
            selected.append(lines[end])
            chars += len(lines[end]) + 1
            end += 1
        text = "\n".join(selected).strip()
        if text:
            chunks.append(
                Chunk(
                    document_path=doc.path,
                    title=doc.title,
                    text=text,
                    ordinal=ordinal,
                    start_line=None if doc.page else cursor + 1,
                    end_line=None if doc.page else end,
                    page=doc.page,
                    tags=doc.tags,
                )
            )
            ordinal += 1
        if end >= len(lines):
            break
        overlap_lines = 0
        overlap_chars = 0
        for line in reversed(selected):
            if overlap_chars >= overlap:
                break
            overlap_chars += len(line) + 1
            overlap_lines += 1
        cursor = max(cursor + 1, end - overlap_lines)
    return chunks


def _source_document(path: Path, absolute: str, text: str, file_type: str) -> SourceDocument:
    cleaned = text.strip()
    tags = tuple(sorted(set(_tags_from_text(cleaned))))
    links = tuple(sorted(set(_links_from_text(cleaned))))
    return SourceDocument(
        absolute,
        path.stem,
        cleaned,
        tags=tags,
        file_type=file_type,
        folder=_folder_label(path),
        links=links,
    )


def _folder_label(path: Path) -> str:
    parent = path.parent.name.strip()
    return parent or "root"


def _links_from_text(text: str) -> list[str]:
    return [match.strip() for match in LINK_PATTERN.findall(text) if match.strip()]


def _tags_from_text(text: str) -> list[str]:
    tags = [match.lower() for match in TAG_PATTERN.findall(text)]
    frontmatter = FRONTMATTER_TAGS.search(text)
    if frontmatter:
        tags.extend(re.findall(r"-\s*([a-zA-Z0-9_\-/]+)", frontmatter.group(1)))
        tags.extend(re.findall(r"tags:\s*\[([^\]]+)\]", frontmatter.group(1)))
    expanded: list[str] = []
    for tag in tags:
        if "," in tag:
            expanded.extend(part.strip().lower() for part in tag.split(",") if part.strip())
        else:
            expanded.append(tag.strip().lower())
    return [tag for tag in expanded if tag]


def _extract_xlsx_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return _extract_xlsx_text_from_xml(path)

    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return _extract_xlsx_text_from_xml(path)
    rows: list[str] = []
    try:
        for worksheet in workbook.worksheets:
            rows.append(f"Sheet: {worksheet.title}")
            for row in worksheet.iter_rows(values_only=True):
                cells = [_format_cell(value) for value in row if value is not None and _format_cell(value)]
                if cells:
                    rows.append("\t".join(cells))
    finally:
        workbook.close()
    return "\n".join(rows)


def _format_cell(value: object) -> str:
    text = str(value).strip()
    return text


def _extract_xlsx_text_from_xml(path: Path) -> str:
    from zipfile import ZipFile

    ns = {
        "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
    }
    with ZipFile(path) as archive:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            for item in root.findall(".//main:si", ns):
                shared_strings.append(" ".join(text.text or "" for text in item.findall(".//main:t", ns)))
        rows: list[str] = []
        for name in sorted(archive.namelist()):
            if not name.startswith("xl/worksheets/sheet") or not name.endswith(".xml"):
                continue
            root = ET.fromstring(archive.read(name))
            for row in root.findall(".//main:row", ns):
                cells: list[str] = []
                for cell in row.findall("main:c", ns):
                    if cell.attrib.get("t") == "inlineStr":
                        text = " ".join(node.text or "" for node in cell.findall(".//main:t", ns)).strip()
                        if text:
                            cells.append(text)
                        continue
                    value = cell.find("main:v", ns)
                    if value is None or value.text is None:
                        continue
                    if cell.attrib.get("t") == "s":
                        index = int(value.text)
                        cells.append(shared_strings[index] if index < len(shared_strings) else value.text)
                    else:
                        cells.append(value.text)
                if cells:
                    rows.append("\t".join(cells))
        return "\n".join(rows)


def ocr_status() -> dict[str, object]:
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    missing = [name for name, path in {"pdftoppm": pdftoppm, "tesseract": tesseract}.items() if not path]
    return {
        "available": not missing,
        "missing": missing,
        "pdftoppm": pdftoppm,
        "tesseract": tesseract,
    }


def _ocr_pdf_page(path: Path, page: int) -> str:
    """OCR one scanned PDF page using local Poppler and Tesseract binaries."""
    status = ocr_status()
    pdftoppm = status["pdftoppm"]
    tesseract = status["tesseract"]
    if not pdftoppm or not tesseract:
        return ""
    with tempfile.TemporaryDirectory(prefix="mnemo-ocr-") as temp:
        output = Path(temp) / "page"
        rendered = subprocess.run(
            [pdftoppm, "-f", str(page), "-l", str(page), "-r", "200", "-png", "-singlefile", str(path), str(output)],
            capture_output=True,
            timeout=90,
            check=False,
        )
        image = output.with_suffix(".png")
        if rendered.returncode or not image.exists():
            return ""
        result = subprocess.run(
            [tesseract, str(image), "stdout", "--psm", "6"],
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        return result.stdout.strip() if result.returncode == 0 else ""
